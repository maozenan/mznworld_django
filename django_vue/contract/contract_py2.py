from docx import Document

document = Document('./contract/contract_work.docx')
def replace_text(old_text,new_text):
    all_paragraphs = document.paragraphs
    for paragraph in all_paragraphs:
        for run in  paragraph.runs:
            run_text = run.text.replace(old_text,new_text)
            run.text = run_text
    all_tables = document.tables
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.replace(old_text,new_text)
                cell.text = cell_text
# list = {
#     "adress": "杭州市萧山区北干街道金城路1号",
#     "party_name": "杭州萧山毛氏集团",
#     "boss": "毛泽楠",
#     "rate": "200",
#     "price_once": "10000",
#     "price_year": "20000",
#     "bank": "萧山建设银行",
#     "banknumber": "18158699969",
#     "number": "91330109MA2G66666",
#     "adress_party": "杭州市萧山区新塘街道董家埭社区",
#     "name": "陈佩瑶",
#     "phone": "18257750308"
# }
def getContract(list):
    count = 1
    for value in list.values():
        old_text = str('X%dX'%count)
        new_text = str(value)
        replace_text(old_text,new_text)
        count = count + 1
    filename = list.get('party_name')
    document.save('./src/file/contract/%s互联网专线合同.docx'%(filename))
    print('ok')